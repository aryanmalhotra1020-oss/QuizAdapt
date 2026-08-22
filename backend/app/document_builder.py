import io
import json
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

OPTION_LETTERS = ['A', 'B', 'C', 'D', 'E', 'F']


def _get_answer_text(question):
    if question.question_type == 'multi_select':
        try:
            correct_list = json.loads(question.correct_answer)
            return ', '.join(correct_list)
        except Exception:
            return question.correct_answer
    return question.correct_answer


def build_docx(subject_name, questions, doc_type):
    doc = Document()
    doc.add_heading(subject_name, level=0)
    doc.add_heading('Answer Key' if doc_type == 'answers' else 'Sample Paper', level=1)

    if doc_type != 'answers':
        info = doc.add_paragraph(f"{len(questions)} questions. Answer all questions.")
        info.runs[0].italic = True

    doc.add_paragraph()

    for i, q in enumerate(questions, 1):
        options = json.loads(q.options) if q.options else None

        if doc_type == 'answers':
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").bold = True
            p.add_run(_get_answer_text(q))
            continue

        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(q.question_text)

        if q.question_type in ('mcq', 'fill_blank') and options:
            for idx, opt in enumerate(options):
                doc.add_paragraph(f"      {OPTION_LETTERS[idx]}. {opt}")
        elif q.question_type == 'multi_select' and options:
            doc.add_paragraph("      (Select all that apply)")
            for opt in options:
                doc.add_paragraph(f"      ☐ {opt}")
        elif q.question_type == 'long_answer':
            for _ in range(3):
                doc.add_paragraph("      " + "_" * 70)

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf(subject_name, questions, doc_type):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('TitleStyle', parent=styles['Title'], fontSize=20, spaceAfter=6)
    subtitle_style = ParagraphStyle('SubtitleStyle', parent=styles['Heading2'], fontSize=14, textColor='#555555', spaceAfter=16)
    question_style = ParagraphStyle('QuestionStyle', parent=styles['Normal'], fontSize=11, spaceBefore=12, spaceAfter=4, leading=15)
    option_style = ParagraphStyle('OptionStyle', parent=styles['Normal'], fontSize=10.5, leftIndent=20, spaceAfter=2, leading=14)
    hint_style = ParagraphStyle('HintStyle', parent=styles['Normal'], fontSize=9, textColor='#888888', spaceAfter=10)

    elements = [Paragraph(subject_name, title_style)]
    elements.append(Paragraph('Answer Key' if doc_type == 'answers' else 'Sample Paper', subtitle_style))

    if doc_type != 'answers':
        elements.append(Paragraph(f"{len(questions)} questions. Answer all questions.", hint_style))

    for i, q in enumerate(questions, 1):
        options = json.loads(q.options) if q.options else None

        if doc_type == 'answers':
            elements.append(Paragraph(f"<b>{i}.</b> {_get_answer_text(q)}", question_style))
            continue

        elements.append(Paragraph(f"<b>{i}.</b> {q.question_text}", question_style))

        if q.question_type in ('mcq', 'fill_blank') and options:
            for idx, opt in enumerate(options):
                elements.append(Paragraph(f"{OPTION_LETTERS[idx]}. {opt}", option_style))
        elif q.question_type == 'multi_select' and options:
            elements.append(Paragraph("<i>(Select all that apply)</i>", option_style))
            for opt in options:
                elements.append(Paragraph(f"&#9744; {opt}", option_style))
        elif q.question_type == 'long_answer':
            elements.append(Spacer(1, 12))
            for _ in range(3):
                elements.append(Paragraph("_" * 90, option_style))
            elements.append(Spacer(1, 6))

    doc.build(elements)
    buffer.seek(0)
    return buffer